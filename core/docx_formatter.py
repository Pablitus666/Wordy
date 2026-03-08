import re
import os
import logging
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from .run_matcher import RunMatcher

class DocxFormatter:
    """ 
    Motor de procesamiento de documentos Word de Nivel Élite.
    Arquitectura con "Super-Inteligencia" Lingüística Corregida.
    """

    def __init__(self):
        self.logger = logging.getLogger("WordFormatterPro.Core")
        self.pattern = None
        self._total_replacements = 0

    def process(self, input_path, words_to_format_str):
        """ Procesa el documento con el motor de búsqueda más avanzado. """
        self.logger.info(f"Iniciando procesamiento de: {input_path}")
        self._total_replacements = 0
        
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Archivo no encontrado: {input_path}")

        # 1. Compilación de la "Super-Regex" (Versión Atómica)
        self.pattern = self._compile_super_regex(words_to_format_str)
        if not self.pattern:
            self.logger.warning("No hay palabras válidas para procesar.")
            return input_path

        self.logger.info(f"Motor de búsqueda activado: {self.pattern.pattern}")

        try:
            doc = Document(input_path)
        except PackageNotFoundError:
            raise Exception("El archivo Word no es válido o está corrupto.")

        # 2. Procesamiento Universal
        for paragraph in self._iter_all_paragraphs(doc):
            if not paragraph.runs or not paragraph.text:
                continue
                
            # Normalización de espacios de Word para el chequeo
            text_to_check = paragraph.text.replace('\xa0', ' ')
            
            if self.pattern.search(text_to_check):
                matcher = RunMatcher(paragraph, self.pattern)
                self._total_replacements += matcher.find_and_format()

        # 3. Guardado con versionado
        output_path = self._generate_output_path(input_path)
        doc.save(output_path)
        
        self.logger.info(f"Finalizado. Coincidencias 'Élite' procesadas: {self._total_replacements}")
        return output_path

    def _compile_super_regex(self, words_str):
        """ 
        Compila una regex que ignora de forma atómica mayúsculas, tildes y género.
        Evita anidamientos de corchetes.
        """
        if not words_str: return None
        
        raw_list = re.split(r'[;,\n\r]+', words_str)
        words = [w.strip() for w in raw_list if w.strip()]
        if not words: return None
        
        words.sort(key=len, reverse=True)
        
        # Mapeo de vocales a sus variantes (Atómico)
        vowel_map = {
            'a': '[aá]', 'á': '[aá]',
            'e': '[eé]', 'é': '[eé]',
            'i': '[ií]', 'í': '[ií]',
            'o': '[oó]', 'ó': '[oó]',
            'u': '[uúü]', 'ú': '[uúü]', 'ü': '[uúü]'
        }

        escaped_parts = []
        for w in words:
            # 1. Escapar caracteres especiales
            escaped_w = re.escape(w)
            
            # 2. Detectar posición del carácter de género (la última o/a antes de signos)
            gender_match = re.search(r'([oaóá])(\\?[^a-zA-ZáéíóúÁÉÍÓÚñÑ]*)$', escaped_w, flags=re.IGNORECASE)
            gender_pos = gender_match.start(1) if gender_match else -1
            
            # 3. Construcción atómica de la regex letra por letra
            final_w = ""
            i = 0
            while i < len(escaped_w):
                char = escaped_w[i]
                lower_char = char.lower()
                
                # Manejo de espacios escapados (Word \xa0)
                if char == '\\' and i+1 < len(escaped_w) and escaped_w[i+1] == ' ':
                    final_w += r'[\s\xa0]+'
                    i += 2
                    continue
                
                # Lógica de Género + Diacríticos
                if i == gender_pos:
                    final_w += '[aáoó]'
                elif lower_char in vowel_map:
                    final_w += vowel_map[lower_char]
                else:
                    final_w += char
                i += 1
            
            escaped_parts.append(final_w)
        
        pattern_str = '|'.join(escaped_parts)
        
        # Lookarounds robustos para límites de palabra en español
        return re.compile(
            rf'(?<![a-zA-Z0-9áéíóúÁÉÍÓÚñÑ])({pattern_str})(?![a-zA-Z0-9áéíóúÁÉÍÓÚñÑ])', 
            re.IGNORECASE
        )

    def _iter_all_paragraphs(self, doc):
        """ Recorre cada sección del documento. """
        yield from doc.paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        for section in doc.sections:
            for container in [section.header, section.first_page_header, section.even_page_header,
                              section.footer, section.first_page_footer, section.even_page_footer]:
                if container and not container.is_linked_to_previous:
                    yield from container.paragraphs

    def _generate_output_path(self, input_path):
        """ Genera nombre de archivo incremental. """
        directory, filename = os.path.split(input_path)
        name, ext = os.path.splitext(filename)
        base = os.path.join(directory, f"{name}_modified{ext}")
        if not os.path.exists(base): return base
        c = 1
        while True:
            out = os.path.join(directory, f"{name}_modified_{c}{ext}")
            if not os.path.exists(out): return out
            c += 1
